#!/usr/bin/env python3
"""Build a clean-reading DOCX from an already-translated PDF.

pdf2docx is used only as a layout parser.  Its Page -> Section -> Column ->
Block model is normalized into a small reading-order IR, then python-docx
creates a fresh flowing document.  We intentionally do not call
pdf2docx.Converter.make_docx(), whose absolute-layout reconstruction is unsuitable
for long Korean textbook prose.
"""

from __future__ import annotations

import argparse
import base64
import binascii
import hashlib
import json
import logging
import math
import os
import re
import statistics
import sys
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable

import fitz
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from pdf2docx import Converter


PDF_MAGIC = b"%PDF-"
ZIP_MAGIC = b"PK"
BLOCK_TEXT = 0
BLOCK_IMAGE = 1
BLOCK_LATTICE_TABLE = 2
BLOCK_STREAM_TABLE = 3
BLOCK_FLOAT_IMAGE = 4
TABLE_TYPES = {BLOCK_LATTICE_TABLE, BLOCK_STREAM_TABLE}
PAGE_NUMBER_RE = re.compile(
    r"^\s*(?:(?:page|쪽)\s*)?(?:\d{1,5}|[ivxlcdm]{1,12})(?:\s*/\s*\d{1,5})?\s*$",
    re.IGNORECASE,
)
HEADING_NUMBER_RE = re.compile(
    r"^\s*(?:제\s*)?(\d+)(?:\.(\d+))?(?:\.(\d+))?(?:\.(\d+))?\s*[.)]?\s+\S+"
)
HEADING_WORD_RE = re.compile(
    r"^\s*(?:제\s*\d+\s*장|\d+\s*장|부록(?:\s+[A-Z가-힣0-9]+)?|appendix\b|chapter\s+\w+\b)",
    re.IGNORECASE,
)
CAPTION_RE = re.compile(
    r"^\s*(?:그림|도표|표|사진|figure|fig\.?|table|photo)\s*[A-Za-z가-힣]*\s*\d+(?:[.\-]\d+)*\s*[:.)-]?",
    re.IGNORECASE,
)
MATH_FONT_RE = re.compile(
    r"(?:math|symbol|stix|cambria|cmr|cmsy|cmmi|latinmodern|computer modern)",
    re.IGNORECASE,
)
KOREAN_RE = re.compile(r"[가-힣]")
LATIN_RE = re.compile(r"[A-Za-z]")
PUNCT_NO_SPACE_BEFORE = set(",.!?:;)]}%‰°℃′″、。，：；！？》〉」』】）")
PUNCT_NO_SPACE_AFTER = set("([{《〈「『【（")
MATH_FUNCTION_RE = re.compile(r"\b(?:sin|cos|tan|sec|csc|cot|log|ln|exp)\b", re.IGNORECASE)
GREEK_RE = re.compile(r"[Α-Ωα-ωϑϕϖϵΣΔΩ]")
FORMULA_INTRO_RE = re.compile(
    r"(?:따라서(?:\s*흔히)?|다음(?:과\s*같이|을\s*얻는다)?|근사값|"
    r"식(?:은|을|으로)?|취하면|미분하면|주어진다|표현된다|계산하면)"
    r"[^.!?]{0,40}[,:，]?\s*$"
)


class LibreOfficeDocxError(RuntimeError):
    """Fail-closed document builder error."""


@dataclass(frozen=True)
class BuildOptions:
    page_size: str = "source"
    page_width_pt: float | None = None
    page_height_pt: float | None = None
    margin_pt: float = 54.0
    font_face: str = "NanumGothic"
    body_size_pt: float = 10.5
    clip_dpi: int = 300
    max_pages: int | None = None
    # Zero-based, end-exclusive page range. This matches PyMuPDF/pdf2docx and
    # lets callers map a source range without an off-by-one translation layer.
    start_page: int = 0
    end_page: int | None = None


def _bbox(value: Any) -> tuple[float, float, float, float]:
    if not isinstance(value, (list, tuple)) or len(value) != 4:
        return (0.0, 0.0, 0.0, 0.0)
    try:
        numbers = tuple(float(v) for v in value)
    except (TypeError, ValueError):
        return (0.0, 0.0, 0.0, 0.0)
    return numbers if all(math.isfinite(v) for v in numbers) else (0.0, 0.0, 0.0, 0.0)


def _bbox_width(value: Any) -> float:
    x0, _, x1, _ = _bbox(value)
    return max(0.0, x1 - x0)


def _bbox_height(value: Any) -> float:
    _, y0, _, y1 = _bbox(value)
    return max(0.0, y1 - y0)


def _clean_text(value: Any) -> str:
    text = str(value or "").replace("\u00ad", "").replace("\ufffd", "")
    text = re.sub(r"[\t\v\f\r ]+", " ", text)
    return text.strip()


def _normalize_repeating_text(value: str) -> str:
    text = _clean_text(value).casefold()
    text = re.sub(r"\s+", " ", text)
    return text[:180]


def _decode_image(value: Any) -> bytes:
    if isinstance(value, bytes):
        data = value
    elif isinstance(value, str) and value:
        try:
            data = base64.b64decode(value.encode("ascii"), validate=True)
        except (ValueError, UnicodeEncodeError, binascii.Error) as exc:
            raise LibreOfficeDocxError("pdf2docx image span has invalid base64 data") from exc
    else:
        raise LibreOfficeDocxError("pdf2docx image span is empty")
    if len(data) < 8:
        raise LibreOfficeDocxError("pdf2docx image span is too small")
    return data


def _span_text(span: dict[str, Any]) -> str:
    return _clean_text(span.get("text", "")) if "image" not in span else ""


def _line_record(line: dict[str, Any]) -> dict[str, Any]:
    text_parts: list[str] = []
    image_parts: list[dict[str, Any]] = []
    fonts: list[str] = []
    sizes: list[float] = []
    bold = False
    italic = False
    for span in line.get("spans", []) or []:
        if not isinstance(span, dict):
            continue
        if span.get("image"):
            image_parts.append(
                {
                    "data": _decode_image(span.get("image")),
                    "bbox": _bbox(span.get("bbox")),
                    "width": float(span.get("width") or _bbox_width(span.get("bbox")) or 1.0),
                    "height": float(span.get("height") or _bbox_height(span.get("bbox")) or 1.0),
                }
            )
            continue
        text = _span_text(span)
        if text:
            text_parts.append(text)
        font = str(span.get("font") or "")
        if font:
            fonts.append(font)
        try:
            size = float(span.get("size") or 0.0)
        except (TypeError, ValueError):
            size = 0.0
        if size > 0 and math.isfinite(size):
            sizes.append(size)
        try:
            flags = int(span.get("flags") or 0)
        except (TypeError, ValueError):
            flags = 0
        bold = bold or bool(flags & 16) or "bold" in font.casefold()
        italic = italic or bool(flags & 2) or "italic" in font.casefold()
    direction = line.get("dir") or (1.0, 0.0)
    try:
        dx, dy = float(direction[0]), float(direction[1])
    except (TypeError, ValueError, IndexError):
        dx, dy = 1.0, 0.0
    rotated = int(line.get("wmode") or 0) != 0 or abs(dx - 1.0) > 0.02 or abs(dy) > 0.02
    return {
        "text": _clean_text(" ".join(text_parts)),
        "images": image_parts,
        "bbox": _bbox(line.get("bbox")),
        "font_size": max(sizes) if sizes else 0.0,
        "fonts": fonts,
        "bold": bold,
        "italic": italic,
        "rotated": rotated,
    }


def _weighted_body_size(pages: list[dict[str, Any]]) -> float:
    samples: list[float] = []

    def visit_blocks(blocks: Iterable[dict[str, Any]]) -> None:
        for block in blocks:
            if not isinstance(block, dict):
                continue
            if int(block.get("type", -1)) in TABLE_TYPES:
                for row in block.get("rows", []) or []:
                    for cell in row.get("cells", []) or []:
                        if cell:
                            visit_blocks(cell.get("blocks", []) or [])
                continue
            for line in block.get("lines", []) or []:
                for span in line.get("spans", []) or []:
                    if not isinstance(span, dict) or span.get("image"):
                        continue
                    text = _span_text(span)
                    try:
                        size = float(span.get("size") or 0.0)
                    except (TypeError, ValueError):
                        continue
                    if text and 4.0 <= size <= 40.0:
                        samples.extend([size] * min(40, max(1, len(text))))

    for page in pages:
        for section in page.get("sections", []) or []:
            for column in section.get("columns", []) or []:
                visit_blocks(column.get("blocks", []) or [])
    return float(statistics.median(samples)) if samples else 10.5


def _join_pair(left: str, right: str) -> str:
    if not left:
        return right
    if not right:
        return left
    if left.endswith(("-", "‐")) and re.search(r"[A-Za-z]$", left[:-1]) and LATIN_RE.match(right):
        return left[:-1] + right
    if right[0] in PUNCT_NO_SPACE_BEFORE or left[-1] in PUNCT_NO_SPACE_AFTER:
        return left + right
    # Korean PDF extraction sometimes leaves a space at every physical line end.
    # One semantic space is preferable to accidental word concatenation.
    return left.rstrip() + " " + right.lstrip()


def _join_lines(lines: list[dict[str, Any]]) -> str:
    result = ""
    for line in lines:
        result = _join_pair(result, line.get("text", ""))
    return _clean_text(result)


def _heading_level(text: str, size: float, body_size: float) -> int | None:
    if not text or len(text) > 180:
        return None
    match = HEADING_NUMBER_RE.match(text)
    if match:
        depth = sum(1 for part in match.groups()[1:] if part is not None)
        # Numbered prose such as "1. 다음을 계산하라" at body size is not a heading.
        if size >= body_size * 1.03 or len(text) <= 80:
            return min(3, 1 + depth)
    if HEADING_WORD_RE.match(text):
        return 1
    if size >= body_size * 1.65 and len(text) <= 120:
        return 1
    if size >= body_size * 1.28 and len(text) <= 120:
        return 2
    if size >= body_size * 1.12 and len(text) <= 100:
        return 3
    return None


def _looks_like_formula(text: str, lines: list[dict[str, Any]], body_size: float = 10.5) -> bool:
    if not text or len(text) > 420:
        return False
    if any(line.get("rotated") for line in lines):
        return True
    fonts = " ".join(font for line in lines for font in line.get("fonts", []))
    if MATH_FONT_RE.search(fonts) and len(text) <= 240:
        return True
    symbol_count = len(re.findall(r"[=≈≠≤≥±∓∑∫∏√∞∂∇∆×÷→←↔^_]", text))
    equation = bool(re.search(r"(?:[A-Za-zΑ-Ωα-ω][\w(){}\[\]′'*/+\- ]{0,30})=", text))
    prose_chars = len(re.findall(r"[A-Za-z가-힣]", text))
    numeric_display = bool(
        re.fullmatch(r"[\d\s.,:;′″'\"°±+−*/×÷()]+", text)
        and len(text) <= 50
    )
    if numeric_display:
        return True
    # PDF math fonts often decode Greek variables to ordinary Unicode while
    # losing operators. Short fragments containing trig functions plus Greek
    # symbols are equations even when '=' is absent (e.g. ``sin Σ cos ( d Σ``).
    if len(text) <= 140 and MATH_FUNCTION_RE.search(text) and GREEK_RE.search(text):
        return True
    if equation and len(text) <= 180:
        return True
    return symbol_count >= 3 and symbol_count / max(1, prose_chars + symbol_count) >= 0.08


def _split_text_block(
    block: dict[str, Any], body_size: float, page_index: int
) -> list[dict[str, Any]]:
    lines = [_line_record(line) for line in block.get("lines", []) or [] if isinstance(line, dict)]
    entries: list[dict[str, Any]] = []
    group: list[dict[str, Any]] = []

    def flush() -> None:
        nonlocal group
        if not group:
            return
        text = _join_lines(group)
        if text:
            size = max((line.get("font_size", 0.0) for line in group), default=body_size)
            level = _heading_level(text, size, body_size)
            numeric_display = bool(re.fullmatch(r"[\d\s.,:;′″'\"°±+−*/×÷()]+", text))
            formula = (
                _looks_like_formula(text, group, body_size)
                and (numeric_display or not HEADING_NUMBER_RE.match(text))
                and not HEADING_WORD_RE.match(text)
            )
            kind = "formula" if formula else ("heading" if level else ("caption" if CAPTION_RE.match(text) else "text"))
            bboxes = [_bbox(line.get("bbox")) for line in group]
            bbox = (
                min(b[0] for b in bboxes),
                min(b[1] for b in bboxes),
                max(b[2] for b in bboxes),
                max(b[3] for b in bboxes),
            )
            entries.append(
                {
                    "kind": kind,
                    "text": text,
                    "bbox": bbox,
                    "page_index": page_index,
                    "font_size": size,
                    "heading_level": level,
                    "bold": any(line.get("bold") for line in group),
                    "italic": any(line.get("italic") for line in group),
                    "rotated": any(line.get("rotated") for line in group),
                }
            )
        group = []

    for line in lines:
        for image in line.get("images", []):
            flush()
            entries.append(
                {
                    "kind": "image",
                    "bbox": image["bbox"],
                    "page_index": page_index,
                    "image": image["data"],
                    "source_width": image["width"],
                    "source_height": image["height"],
                }
            )
        if not line.get("text"):
            continue
        line_level = _heading_level(line["text"], line.get("font_size", 0.0), body_size)
        line_caption = bool(CAPTION_RE.match(line["text"]))
        if group:
            previous = group[-1]
            prev_level = _heading_level(previous["text"], previous.get("font_size", 0.0), body_size)
            prev_caption = bool(CAPTION_RE.match(previous["text"]))
            gap = line["bbox"][1] - previous["bbox"][3]
            size_a = max(1.0, previous.get("font_size", body_size))
            size_b = max(1.0, line.get("font_size", body_size))
            role_changed = (line_level is not None) != (prev_level is not None) or line_caption != prev_caption
            size_changed = max(size_a, size_b) / min(size_a, size_b) > 1.22
            if role_changed or size_changed or gap > max(8.0, body_size * 1.45):
                flush()
        group.append(line)
    flush()
    return entries


def _cell_text(cell: dict[str, Any] | None) -> str:
    if not cell:
        return ""
    parts: list[str] = []
    for block in cell.get("blocks", []) or []:
        if int(block.get("type", -1)) in TABLE_TYPES:
            return ""
        for line in block.get("lines", []) or []:
            record = _line_record(line)
            if record.get("images"):
                return ""
            if record.get("text"):
                parts.append(record["text"])
    return _join_lines([{"text": part} for part in parts])


def _table_has_rotation(block: dict[str, Any]) -> bool:
    for row in block.get("rows", []) or []:
        for cell in row.get("cells", []) or []:
            if not cell:
                continue
            for child in cell.get("blocks", []) or []:
                for line in child.get("lines", []) or []:
                    if _line_record(line).get("rotated"):
                        return True
    return False


def table_is_safe(block: dict[str, Any]) -> tuple[bool, list[list[str]]]:
    rows = block.get("rows", []) or []
    if not 1 <= len(rows) <= 28:
        return False, []
    matrices: list[list[str]] = []
    column_count: int | None = None
    total_chars = 0
    for row in rows:
        cells = row.get("cells", []) or []
        if column_count is None:
            column_count = len(cells)
        if not column_count or column_count > 8 or len(cells) != column_count:
            return False, []
        values: list[str] = []
        for cell in cells:
            if not cell:
                return False, []
            merged = tuple(cell.get("merged_cells") or (1, 1))
            if merged != (1, 1):
                return False, []
            blocks = cell.get("blocks", []) or []
            if any(int(child.get("type", -1)) in TABLE_TYPES for child in blocks):
                return False, []
            value = _cell_text(cell)
            if not value and any(
                span.get("image")
                for child in blocks
                for line in child.get("lines", []) or []
                for span in line.get("spans", []) or []
                if isinstance(span, dict)
            ):
                return False, []
            cell_lines = [
                _line_record(line)
                for child in blocks
                for line in child.get("lines", []) or []
                if isinstance(line, dict)
            ]
            # Formula boxes often arrive as visually simple 1x2/1x3 tables (the
            # last cell is an equation number). Rebuilding those as Word cells
            # loses glyph positioning, so preserve the whole source region.
            if len(rows) <= 2 and value and _looks_like_formula(value, cell_lines):
                return False, []
            values.append(value)
            total_chars += len(value)
        matrices.append(values)
    if total_chars > 4500 or _table_has_rotation(block):
        return False, []
    # A one-cell "table" is usually a framed callout or equation, not tabular data.
    if len(matrices) == 1 and (column_count or 0) == 1:
        return False, []
    return True, matrices


def _table_text_matrix(block: dict[str, Any]) -> list[list[str]]:
    matrix: list[list[str]] = []
    for row in block.get("rows", []) or []:
        values = [_cell_text(cell) for cell in row.get("cells", []) or [] if cell]
        if values:
            matrix.append(values)
    return matrix


def _cell_bbox_text(cell: dict[str, Any] | None) -> tuple[tuple[float, float, float, float], str] | None:
    if not cell:
        return None
    return _bbox(cell.get("bbox")), _cell_text(cell)


def _union_bboxes(values: Iterable[Any]) -> tuple[float, float, float, float]:
    boxes = [_bbox(value) for value in values]
    boxes = [box for box in boxes if box[2] > box[0] and box[3] > box[1]]
    if not boxes:
        return (0.0, 0.0, 0.0, 0.0)
    return (
        min(box[0] for box in boxes),
        min(box[1] for box in boxes),
        max(box[2] for box in boxes),
        max(box[3] for box in boxes),
    )


def _split_mixed_formula_caption_table(
    block: dict[str, Any], page_index: int
) -> list[dict[str, Any]] | None:
    cells: list[tuple[tuple[float, float, float, float], str]] = []
    for row in block.get("rows", []) or []:
        for cell in row.get("cells", []) or []:
            value = _cell_bbox_text(cell)
            if value and value[1]:
                cells.append(value)
    caption_cells = [value for value in cells if CAPTION_RE.match(value[1])]
    if not caption_cells:
        return None
    caption_x0 = min(value[0][0] for value in caption_cells)
    left = [value for value in cells if value[0][2] <= caption_x0 + 0.5]
    right = [value for value in cells if value[0][0] >= caption_x0 - 0.5]
    if not left or not right:
        return None
    left_text = " ".join(value[1] for value in left)
    if not re.search(r"[=≈≠≤≥±∓∑∫√]|\(\d+\.\d+\)", left_text):
        return None
    caption_text = _join_lines([{"text": value[1]} for value in right])
    return [
        {
            "kind": "table_clip",
            "bbox": _union_bboxes(value[0] for value in left),
            "page_index": page_index,
            "rows": None,
            "source_rows": [[value[1] for value in left]],
        },
        {
            "kind": "caption_clip",
            "bbox": _union_bboxes(value[0] for value in right),
            "page_index": page_index,
            "text": caption_text,
        },
    ]


def _image_entry(block: dict[str, Any], page_index: int) -> dict[str, Any]:
    return {
        "kind": "image",
        "bbox": _bbox(block.get("bbox")),
        "page_index": page_index,
        "image": _decode_image(block.get("image")),
        "source_width": float(block.get("width") or _bbox_width(block.get("bbox")) or 1.0),
        "source_height": float(block.get("height") or _bbox_height(block.get("bbox")) or 1.0),
    }


def _block_entries(block: dict[str, Any], body_size: float, page_index: int) -> list[dict[str, Any]]:
    block_type = int(block.get("type", -1))
    if block_type == BLOCK_TEXT:
        return _split_text_block(block, body_size, page_index)
    if block_type in (BLOCK_IMAGE, BLOCK_FLOAT_IMAGE):
        return [_image_entry(block, page_index)]
    if block_type in TABLE_TYPES:
        mixed = _split_mixed_formula_caption_table(block, page_index)
        if mixed:
            return mixed
        safe, matrix = table_is_safe(block)
        source_matrix = matrix if safe else _table_text_matrix(block)
        return [
            {
                "kind": "table" if safe else "table_clip",
                "bbox": _bbox(block.get("bbox")),
                "page_index": page_index,
                "rows": matrix if safe else None,
                "source_rows": source_matrix,
            }
        ]
    return []


def _insert_floats(columns: list[dict[str, Any]], floats: list[dict[str, Any]]) -> None:
    if not columns:
        return
    for item in floats:
        center_x = (item["bbox"][0] + item["bbox"][2]) / 2.0
        target = min(
            columns,
            key=lambda col: abs(center_x - (col["bbox"][0] + col["bbox"][2]) / 2.0),
        )
        target["blocks"].append(item)
    for column in columns:
        column["blocks"].sort(key=lambda item: (item["bbox"][1], item["bbox"][0]))


def _is_formula_intro(text: str) -> bool:
    value = _clean_text(text)
    return bool(value.endswith((",", ":", "，")) or FORMULA_INTRO_RE.search(value))


def _repair_detached_bottom_visuals(ir: dict[str, Any]) -> dict[str, int]:
    repaired = 0
    for page in ir["pages"]:
        sections = page["sections"]
        index = 0
        while index + 1 < len(sections):
            current = sections[index]
            following = sections[index + 1]
            if len(current["columns"]) != 2 or len(following["columns"]) != 1:
                index += 1
                continue
            left, right = current["columns"]
            next_blocks = following["columns"][0]["blocks"]
            if not left["blocks"] or not right["blocks"] or not next_blocks:
                index += 1
                continue
            intro = left["blocks"][-1]
            if intro.get("kind") != "text" or not _is_formula_intro(intro.get("text", "")):
                index += 1
                continue
            if right["blocks"][0].get("kind") in ("formula", "table_clip"):
                index += 1
                continue
            if any(block.get("kind") not in ("formula", "table_clip", "caption_clip") for block in next_blocks):
                index += 1
                continue
            formula_blocks = [block for block in next_blocks if block.get("kind") in ("formula", "table_clip")]
            caption_blocks = [block for block in next_blocks if block.get("kind") == "caption_clip"]
            if not formula_blocks:
                index += 1
                continue
            left["blocks"].extend(formula_blocks)
            right["blocks"].extend(caption_blocks)
            sections.pop(index + 1)
            repaired += len(formula_blocks)
            # Keep index on the current section in case another detached visual
            # section immediately follows it.
    return {"detached_formula_clips_reanchored": repaired}


def _vertical_gap(a: Any, b: Any) -> float:
    first, second = _bbox(a), _bbox(b)
    return max(0.0, first[1] - second[3], second[1] - first[3])


def _group_figures_and_captions(ir: dict[str, Any]) -> dict[str, int]:
    grouped = 0
    clipped = 0
    joined_fragments = 0
    for page in ir["pages"]:
        for section in page["sections"]:
            for column in section["columns"]:
                blocks = column["blocks"]
                cursor = 0
                while cursor < len(blocks):
                    if blocks[cursor].get("kind") != "image":
                        cursor += 1
                        continue
                    image_index = cursor
                    candidates = []
                    for candidate_index in range(max(0, image_index - 2), min(len(blocks), image_index + 3)):
                        candidate = blocks[candidate_index]
                        if candidate.get("kind") not in ("caption", "caption_clip"):
                            continue
                        if _vertical_gap(candidate["bbox"], blocks[image_index]["bbox"]) <= 20.0:
                            candidates.append((abs(candidate_index - image_index), candidate_index))
                    if not candidates:
                        cursor += 1
                        continue
                    caption_index = min(candidates)[1]
                    caption = blocks[caption_index]
                    continuation_indexes: list[int] = []
                    caption_end = caption["bbox"][3]
                    scan = max(image_index, caption_index) + 1
                    while scan < len(blocks) and len(continuation_indexes) < 4:
                        candidate = blocks[scan]
                        if candidate.get("kind") not in ("text", "formula"):
                            break
                        if candidate["bbox"][1] - caption_end > 16.0:
                            break
                        continuation_indexes.append(scan)
                        caption_end = max(caption_end, candidate["bbox"][3])
                        scan += 1
                    fragments = [caption] + [blocks[position] for position in continuation_indexes]
                    use_clip = caption.get("kind") == "caption_clip" or any(
                        fragment.get("kind") == "formula" for fragment in fragments
                    )
                    caption_bbox = _union_bboxes(fragment["bbox"] for fragment in fragments)
                    caption_text = _join_lines(
                        [{"text": fragment.get("text", "")} for fragment in fragments]
                    )
                    image = blocks[image_index]
                    figure = {
                        "kind": "figure",
                        "bbox": _union_bboxes((image["bbox"], caption_bbox)),
                        "page_index": image["page_index"],
                        "image": image["image"],
                        "image_bbox": image["bbox"],
                        "source_width": image.get("source_width"),
                        "source_height": image.get("source_height"),
                        "caption_mode": "clip" if use_clip else "text",
                        "caption_bbox": caption_bbox,
                        "caption_text": caption_text,
                    }
                    removed = {image_index, caption_index, *continuation_indexes}
                    insertion = min(removed)
                    blocks[:] = [block for position, block in enumerate(blocks) if position not in removed]
                    blocks.insert(insertion, figure)
                    grouped += 1
                    joined_fragments += len(continuation_indexes)
                    if use_clip:
                        clipped += 1
                    cursor = insertion + 1
    return {
        "atomic_figures": grouped,
        "caption_fragments_joined": joined_fragments,
        "caption_clips": clipped,
    }


def _suppress_running_furniture(ir: dict[str, Any]) -> dict[str, int]:
    pages = ir["pages"]
    threshold = 2 if len(pages) <= 4 else 3
    text_counts: dict[str, int] = {}
    image_counts: dict[str, int] = {}

    def edge(item: dict[str, Any], page: dict[str, Any]) -> bool:
        y0, y1 = item["bbox"][1], item["bbox"][3]
        return y1 <= page["height"] * 0.12 or y0 >= page["height"] * 0.88

    for page in pages:
        seen_text: set[str] = set()
        seen_images: set[str] = set()
        for section in page["sections"]:
            for column in section["columns"]:
                for item in column["blocks"]:
                    if not edge(item, page):
                        continue
                    if item["kind"] in ("text", "heading", "caption"):
                        key = _normalize_repeating_text(item.get("text", ""))
                        if key and len(key) <= 160:
                            seen_text.add(key)
                    elif item["kind"] == "image":
                        seen_images.add(hashlib.sha256(item["image"]).hexdigest())
        for key in seen_text:
            text_counts[key] = text_counts.get(key, 0) + 1
        for key in seen_images:
            image_counts[key] = image_counts.get(key, 0) + 1

    removed_headers = 0
    removed_numbers = 0
    removed_images = 0
    removed_header_tables = 0
    for page in pages:
        explicit = {
            _normalize_repeating_text(page.get("header", "")),
            _normalize_repeating_text(page.get("footer", "")),
        }
        for section in page["sections"]:
            for column in section["columns"]:
                kept = []
                for item in column["blocks"]:
                    if not edge(item, page):
                        kept.append(item)
                        continue
                    if item["kind"] in ("text", "heading", "caption", "formula"):
                        text = item.get("text", "")
                        key = _normalize_repeating_text(text)
                        if PAGE_NUMBER_RE.match(text):
                            removed_numbers += 1
                            continue
                        if item["kind"] != "formula" and key and (key in explicit or text_counts.get(key, 0) >= threshold):
                            removed_headers += 1
                            continue
                    elif item["kind"] in ("table", "table_clip"):
                        rows = item.get("source_rows") or item.get("rows") or []
                        flattened = [value for row in rows for value in row if _clean_text(value)]
                        has_standalone_page = any(PAGE_NUMBER_RE.match(value) for value in flattened)
                        total_text = " ".join(flattened)
                        # Textbook running furniture is commonly a one-row box with
                        # chapter/section title in one cell and the page number in the
                        # other. Parenthesized equation numbers do not match the page
                        # regex, so formula boxes at the page edge are retained.
                        if len(rows) == 1 and len(flattened) <= 3 and has_standalone_page and len(total_text) <= 120:
                            removed_header_tables += 1
                            continue
                    elif item["kind"] == "image":
                        key = hashlib.sha256(item["image"]).hexdigest()
                        if image_counts.get(key, 0) >= threshold:
                            removed_images += 1
                            continue
                    kept.append(item)
                column["blocks"] = kept
    return {
        "repeating_text_removed": removed_headers,
        "page_numbers_removed": removed_numbers,
        "repeating_images_removed": removed_images,
        "running_header_tables_removed": removed_header_tables,
    }


def _serialized_to_ir(serialized: list[dict[str, Any]]) -> tuple[dict[str, Any], dict[str, int]]:
    body_size = _weighted_body_size(serialized)
    pages_ir: list[dict[str, Any]] = []
    for raw_page in serialized:
        page_index = int(raw_page.get("id", len(pages_ir)))
        sections_ir: list[dict[str, Any]] = []
        for raw_section in sorted(
            raw_page.get("sections", []) or [], key=lambda value: (_bbox(value.get("bbox"))[1], _bbox(value.get("bbox"))[0])
        ):
            columns_ir: list[dict[str, Any]] = []
            for raw_column in sorted(
                raw_section.get("columns", []) or [], key=lambda value: _bbox(value.get("bbox"))[0]
            ):
                entries: list[dict[str, Any]] = []
                for block in sorted(
                    raw_column.get("blocks", []) or [], key=lambda value: (_bbox(value.get("bbox"))[1], _bbox(value.get("bbox"))[0])
                ):
                    entries.extend(_block_entries(block, body_size, page_index))
                entries.sort(key=lambda item: (item["bbox"][1], item["bbox"][0]))
                columns_ir.append({"bbox": _bbox(raw_column.get("bbox")), "blocks": entries})
            sections_ir.append(
                {
                    "bbox": _bbox(raw_section.get("bbox")),
                    "columns": columns_ir,
                    "num_cols": len(columns_ir),
                }
            )
        floats = []
        for raw_float in raw_page.get("floats", []) or []:
            try:
                floats.append(_image_entry(raw_float, page_index))
            except LibreOfficeDocxError:
                raise
        all_columns = [column for section in sections_ir for column in section["columns"]]
        _insert_floats(all_columns, floats)
        pages_ir.append(
            {
                "index": page_index,
                "width": float(raw_page.get("width") or 595.0),
                "height": float(raw_page.get("height") or 842.0),
                "header": raw_page.get("header", ""),
                "footer": raw_page.get("footer", ""),
                "sections": sections_ir,
            }
        )
    ir = {"body_font_size": body_size, "pages": pages_ir}
    repair = _repair_detached_bottom_visuals(ir)
    figures = _group_figures_and_captions(ir)
    removal = _suppress_running_furniture(ir)
    return ir, {**repair, **figures, **removal}


def parse_pdf_to_ir(pdf_path: str | os.PathLike[str], options: BuildOptions) -> tuple[dict[str, Any], dict[str, Any]]:
    source = Path(pdf_path)
    data = source.read_bytes()
    if not data.startswith(PDF_MAGIC):
        raise LibreOfficeDocxError("input is not a PDF")
    try:
        probe = fitz.open(source)
    except Exception as exc:
        raise LibreOfficeDocxError(f"cannot open translated PDF: {exc}") from exc
    try:
        if probe.needs_pass:
            raise LibreOfficeDocxError("encrypted PDFs are not supported")
        total_pages = probe.page_count
    finally:
        probe.close()
    if total_pages <= 0:
        raise LibreOfficeDocxError("translated PDF has no pages")
    start_page = options.start_page
    end_page = options.end_page if options.end_page is not None else total_pages
    if start_page < 0 or end_page <= start_page or end_page > total_pages:
        raise LibreOfficeDocxError(
            f"invalid zero-based page range [{start_page}, {end_page}) for {total_pages} pages"
        )
    if options.max_pages:
        end_page = min(end_page, start_page + options.max_pages)
    expected_pages = end_page - start_page

    converter = Converter(str(source))
    try:
        settings = converter.default_settings
        settings.update(
            {
                "ignore_page_error": False,
                "raw_exceptions": True,
                "multi_processing": False,
                "delete_end_line_hyphen": False,
                "extract_stream_table": False,
                "parse_lattice_table": True,
                "parse_stream_table": True,
                "list_not_table": True,
            }
        )
        converter.parse(start=start_page, end=end_page, **settings)
        serialized = [
            page
            for page in converter.pages.store()
            if start_page <= int(page.get("id", -1)) < end_page
        ]
    except Exception as exc:
        raise LibreOfficeDocxError(f"pdf2docx layout parsing failed: {exc}") from exc
    finally:
        converter.close()
    if len(serialized) != expected_pages:
        raise LibreOfficeDocxError(
            f"pdf2docx parsed {len(serialized)} pages but {expected_pages} were required"
        )
    ir, removal = _serialized_to_ir(serialized)
    return ir, {
        "source_pages": total_pages,
        "processed_pages": expected_pages,
        "start_page_zero_based": start_page,
        "end_page_zero_based_exclusive": end_page,
        **removal,
    }


def _set_run_font(run: Any, face: str, size: float, *, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = face
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), face)


def _configure_styles(document: Document, options: BuildOptions) -> None:
    palette = RGBColor(31, 55, 77)
    normal = document.styles["Normal"]
    normal.font.name = options.font_face
    normal.font.size = Pt(options.body_size_pt)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), options.font_face)
    normal.paragraph_format.space_after = Pt(5.0)
    normal.paragraph_format.line_spacing = 1.22
    for level, size in ((1, 18.0), (2, 15.0), (3, 12.5)):
        style = document.styles[f"Heading {level}"]
        style.font.name = options.font_face
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = palette
        style._element.rPr.rFonts.set(qn("w:eastAsia"), options.font_face)
        style.paragraph_format.space_before = Pt(14.0 if level == 1 else 10.0)
        style.paragraph_format.space_after = Pt(5.0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def _resolve_page_geometry(ir: dict[str, Any], options: BuildOptions) -> tuple[float, float]:
    if options.page_width_pt is not None or options.page_height_pt is not None:
        if options.page_width_pt is None or options.page_height_pt is None:
            raise LibreOfficeDocxError("page width and height must be provided together")
        width, height = options.page_width_pt, options.page_height_pt
    elif options.page_size == "a4":
        width, height = 595.276, 841.89
    elif options.page_size == "letter":
        width, height = 612.0, 792.0
    elif options.page_size == "source":
        first = ir["pages"][0]
        width, height = float(first["width"]), float(first["height"])
    else:
        raise LibreOfficeDocxError(f"unsupported page size: {options.page_size}")
    if not (200.0 <= width <= 2000.0 and 200.0 <= height <= 2000.0):
        raise LibreOfficeDocxError(f"unsafe page size: {width}x{height} pt")
    if options.margin_pt * 2 >= min(width, height) - 72.0:
        raise LibreOfficeDocxError("page margins leave no readable content width")
    return width, height


def _set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Pt(width / 20.0)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def _table_widths(rows: list[list[str]], usable_width_dxa: int) -> list[int]:
    columns = len(rows[0])
    weights = []
    for col in range(columns):
        longest = max((len(re.sub(r"\s+", " ", row[col])) for row in rows), default=1)
        weights.append(min(40.0, max(5.0, math.sqrt(max(1, longest)) * 4.0)))
    total_weight = sum(weights)
    widths = [max(900, round(usable_width_dxa * weight / total_weight)) for weight in weights]
    scale = usable_width_dxa / sum(widths)
    widths = [max(600, round(width * scale)) for width in widths]
    widths[-1] += usable_width_dxa - sum(widths)
    return widths


def _add_real_table(document: Document, rows: list[list[str]], usable_width_pt: float, options: BuildOptions) -> None:
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = _table_widths(rows, int(round(usable_width_pt * 20)))
    _set_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)
            if row_index == 0:
                _set_cell_shading(cell, "DCE6F1")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) <= 20 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            _set_run_font(run, options.font_face, max(8.5, options.body_size_pt - 1.0), bold=row_index == 0)
    _set_repeat_table_header(table.rows[0])
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _clip_pdf_region(pdf: fitz.Document, item: dict[str, Any], dpi: int) -> bytes:
    page_index = int(item["page_index"])
    if page_index < 0 or page_index >= pdf.page_count:
        raise LibreOfficeDocxError(f"clip references invalid source page {page_index + 1}")
    page = pdf[page_index]
    clip = fitz.Rect(item["bbox"]) & page.rect
    if clip.is_empty or clip.width < 1.0 or clip.height < 1.0:
        raise LibreOfficeDocxError(f"clip has invalid bbox on source page {page_index + 1}")
    pixmap = page.get_pixmap(dpi=dpi, clip=clip, alpha=False)
    data = pixmap.tobytes("png")
    if len(data) < 32:
        raise LibreOfficeDocxError(f"clip render is empty on source page {page_index + 1}")
    return data


def _add_centered_image(
    document: Document,
    data: bytes,
    width_pt: float,
    usable_width_pt: float,
    options: BuildOptions,
    alt_text: str,
    *,
    keep_with_next: bool = False,
) -> Any:
    width = min(usable_width_pt, max(36.0, width_pt))
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.keep_with_next = keep_with_next
    try:
        shape = paragraph.add_run().add_picture(BytesIO(data), width=Pt(width))
    except Exception as exc:
        raise LibreOfficeDocxError(f"cannot insert image into DOCX: {exc}") from exc
    shape._inline.docPr.set("descr", alt_text[:255])
    return paragraph


def _add_caption_paragraph(document: Document, text: str, options: BuildOptions) -> Any:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_together = True
    size = max(8.5, options.body_size_pt - 1.0)
    run = paragraph.add_run(text)
    _set_run_font(run, options.font_face, size, italic=True)
    run.font.color.rgb = RGBColor(89, 89, 89)
    return paragraph


def _iter_reading_order(ir: dict[str, Any]) -> Iterable[dict[str, Any]]:
    for page in ir["pages"]:
        for section in page["sections"]:
            # Columns are intentionally exhausted left-to-right. Sorting all blocks by
            # y would interleave two-column prose incorrectly.
            for column in section["columns"]:
                yield from column["blocks"]


def build_docx(
    pdf_path: str | os.PathLike[str],
    output_path: str | os.PathLike[str],
    options: BuildOptions,
) -> dict[str, Any]:
    ir, parse_stats = parse_pdf_to_ir(pdf_path, options)
    width_pt, height_pt = _resolve_page_geometry(ir, options)
    usable_width_pt = width_pt - options.margin_pt * 2

    document = Document()
    _configure_styles(document, options)
    section = document.sections[0]
    section.page_width = Pt(width_pt)
    section.page_height = Pt(height_pt)
    section.orientation = WD_ORIENT.LANDSCAPE if width_pt > height_pt else WD_ORIENT.PORTRAIT
    section.top_margin = Pt(options.margin_pt)
    section.bottom_margin = Pt(options.margin_pt)
    section.left_margin = Pt(options.margin_pt)
    section.right_margin = Pt(options.margin_pt)
    document.core_properties.title = "Quilo PDF translation - LibreOffice clean reading edition"
    document.core_properties.author = "Quilo"
    document.core_properties.subject = "Translated PDF clean-reading reflow"

    counts = {
        "paragraphs": 0,
        "headings": 0,
        "captions": 0,
        "images": 0,
        "formula_clips": 0,
        "caption_clips": 0,
        "real_tables": 0,
        "table_clips": 0,
        "duplicate_ir_images_skipped": 0,
    }
    # pdf2docx can expose the same underlying image as both an inline and a
    # floating block. Suppress only that exact duplicate IR occurrence. The same
    # image at another bbox or on another page is legitimate content and must be
    # preserved (logos, repeated diagrams, legends, etc.).
    seen_image_occurrences: set[tuple[int, tuple[float, float, float, float], str]] = set()
    pdf = fitz.open(pdf_path)
    try:
        reading_items = list(_iter_reading_order(ir))
        for item_index, item in enumerate(reading_items):
            kind = item["kind"]
            next_kind = reading_items[item_index + 1]["kind"] if item_index + 1 < len(reading_items) else None
            if kind == "heading":
                level = int(item.get("heading_level") or 2)
                paragraph = document.add_paragraph(style=f"Heading {min(3, max(1, level))}")
                run = paragraph.add_run(item["text"])
                _set_run_font(run, options.font_face, document.styles[f"Heading {min(3, max(1, level))}"].font.size.pt, bold=True)
                counts["headings"] += 1
            elif kind in ("text", "caption"):
                if kind == "caption":
                    _add_caption_paragraph(document, item["text"], options)
                    counts["captions"] += 1
                else:
                    paragraph = document.add_paragraph()
                    paragraph.paragraph_format.widow_control = True
                    paragraph.paragraph_format.keep_together = False
                    paragraph.paragraph_format.first_line_indent = Pt(options.body_size_pt)
                    if next_kind in ("formula", "table_clip") and _is_formula_intro(item.get("text", "")):
                        paragraph.paragraph_format.keep_with_next = True
                    run = paragraph.add_run(item["text"])
                    _set_run_font(
                        run,
                        options.font_face,
                        options.body_size_pt,
                        bold=bool(item.get("bold")),
                        italic=bool(item.get("italic")),
                    )
                    counts["paragraphs"] += 1
            elif kind == "image":
                digest = hashlib.sha256(item["image"]).hexdigest()
                occurrence = (
                    int(item["page_index"]),
                    tuple(round(value, 2) for value in item["bbox"]),
                    digest,
                )
                if occurrence in seen_image_occurrences:
                    counts["duplicate_ir_images_skipped"] += 1
                    continue
                seen_image_occurrences.add(occurrence)
                display_width = _bbox_width(item["bbox"]) or float(item.get("source_width") or usable_width_pt)
                _add_centered_image(
                    document,
                    item["image"],
                    display_width,
                    usable_width_pt,
                    options,
                    f"Source image from translated PDF page {item['page_index'] + 1}",
                )
                counts["images"] += 1
            elif kind == "figure":
                image_bbox = item["image_bbox"]
                digest = hashlib.sha256(item["image"]).hexdigest()
                occurrence = (
                    int(item["page_index"]),
                    tuple(round(value, 2) for value in image_bbox),
                    digest,
                )
                if occurrence in seen_image_occurrences:
                    counts["duplicate_ir_images_skipped"] += 1
                    continue
                seen_image_occurrences.add(occurrence)
                _add_centered_image(
                    document,
                    item["image"],
                    _bbox_width(image_bbox) or float(item.get("source_width") or usable_width_pt),
                    usable_width_pt,
                    options,
                    f"Figure from translated PDF page {item['page_index'] + 1}",
                    keep_with_next=True,
                )
                counts["images"] += 1
                if item["caption_mode"] == "clip":
                    caption_data = _clip_pdf_region(
                        pdf,
                        {**item, "bbox": item["caption_bbox"]},
                        options.clip_dpi,
                    )
                    _add_centered_image(
                        document,
                        caption_data,
                        _bbox_width(item["caption_bbox"]),
                        usable_width_pt,
                        options,
                        f"Figure caption preserved from translated PDF page {item['page_index'] + 1}",
                    )
                    counts["caption_clips"] += 1
                else:
                    _add_caption_paragraph(document, item["caption_text"], options)
                    counts["captions"] += 1
            elif kind == "table":
                rows = item.get("rows") or []
                if not rows or not rows[0]:
                    raise LibreOfficeDocxError("safe table has no cells")
                _add_real_table(document, rows, usable_width_pt, options)
                counts["real_tables"] += 1
            elif kind in ("formula", "table_clip", "caption_clip"):
                data = _clip_pdf_region(pdf, item, options.clip_dpi)
                display_width = _bbox_width(item["bbox"])
                _add_centered_image(
                    document,
                    data,
                    display_width,
                    usable_width_pt,
                    options,
                    ("Formula" if kind == "formula" else "Caption" if kind == "caption_clip" else "Complex table")
                    + f" preserved from translated PDF page {item['page_index'] + 1}",
                )
                counts[
                    "formula_clips" if kind == "formula" else "caption_clips" if kind == "caption_clip" else "table_clips"
                ] += 1
            else:
                raise LibreOfficeDocxError(f"unsupported IR block kind: {kind}")
    finally:
        pdf.close()

    if sum(counts[key] for key in ("paragraphs", "headings", "captions", "images", "formula_clips", "caption_clips", "real_tables", "table_clips")) == 0:
        raise LibreOfficeDocxError("translated PDF produced an empty clean-reading document")
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    data = output.read_bytes()
    if not data.startswith(ZIP_MAGIC) or len(data) < 1000:
        raise LibreOfficeDocxError("generated DOCX is missing, empty, or invalid")
    return {
        "builder": "pdf2docx-clean-reading-v1",
        "page_size": {"width_pt": width_pt, "height_pt": height_pt, "mode": options.page_size},
        "font_face": options.font_face,
        "body_size_pt": options.body_size_pt,
        "clip_dpi": options.clip_dpi,
        "docx_bytes": len(data),
        **parse_stats,
        **counts,
    }


def _parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_pdf")
    parser.add_argument("output_docx")
    parser.add_argument("--metadata", required=True)
    parser.add_argument("--page-size", choices=("source", "a4", "letter"), default="source")
    parser.add_argument("--page-width-pt", type=float)
    parser.add_argument("--page-height-pt", type=float)
    parser.add_argument("--margin-pt", type=float, default=54.0)
    parser.add_argument("--font-face", default="NanumGothic")
    parser.add_argument("--body-size-pt", type=float, default=10.5)
    parser.add_argument("--clip-dpi", type=int, default=300)
    parser.add_argument("--max-pages", type=int)
    parser.add_argument("--start-page", type=int, default=0, help="zero-based first page")
    parser.add_argument("--end-page", type=int, help="zero-based exclusive end page")
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = _parse_args(argv or sys.argv[1:])
    logging.getLogger().setLevel(logging.WARNING)
    try:
        if not 6.0 <= args.body_size_pt <= 24.0:
            raise LibreOfficeDocxError("body size must be between 6 and 24 pt")
        if not 36.0 <= args.margin_pt <= 144.0:
            raise LibreOfficeDocxError("margin must be between 36 and 144 pt")
        if not 144 <= args.clip_dpi <= 600:
            raise LibreOfficeDocxError("clip DPI must be between 144 and 600")
        if args.max_pages is not None and args.max_pages <= 0:
            raise LibreOfficeDocxError("max pages must be positive")
        if args.start_page < 0:
            raise LibreOfficeDocxError("start page must be zero or greater")
        if args.end_page is not None and args.end_page <= args.start_page:
            raise LibreOfficeDocxError("end page must be greater than start page")
        options = BuildOptions(
            page_size=args.page_size,
            page_width_pt=args.page_width_pt,
            page_height_pt=args.page_height_pt,
            margin_pt=args.margin_pt,
            font_face=_clean_text(args.font_face) or "NanumGothic",
            body_size_pt=args.body_size_pt,
            clip_dpi=args.clip_dpi,
            max_pages=args.max_pages,
            start_page=args.start_page,
            end_page=args.end_page,
        )
        metadata = build_docx(args.input_pdf, args.output_docx, options)
        metadata_path = Path(args.metadata)
        metadata_path.parent.mkdir(parents=True, exist_ok=True)
        metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")
        return 0
    except Exception as exc:
        print(f"LibreOffice clean-reading DOCX failed: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
