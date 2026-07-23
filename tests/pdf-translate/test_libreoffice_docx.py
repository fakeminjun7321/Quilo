import importlib.util
import os
import shutil
import sys
import tempfile
import unittest
from pathlib import Path

import fitz
from docx import Document


ROOT = Path(__file__).resolve().parents[2]
SCRIPT = ROOT / "lib" / "pipelines" / "pdf-translate" / "libreoffice-docx.py"
SPEC = importlib.util.spec_from_file_location("quilo_libreoffice_docx", SCRIPT)
MODULE = importlib.util.module_from_spec(SPEC)
sys.modules[SPEC.name] = MODULE
SPEC.loader.exec_module(MODULE)


def text_cell(value, *, rotated=False):
    return {
        "merged_cells": [1, 1],
        "blocks": [
            {
                "type": 0,
                "lines": [
                    {
                        "bbox": [0, 0, 30, 10],
                        "dir": [0, 1] if rotated else [1, 0],
                        "spans": [
                            {
                                "bbox": [0, 0, 30, 10],
                                "text": value,
                                "font": "Pretendard",
                                "size": 10,
                            }
                        ],
                    }
                ],
            }
        ],
    }


class LibreOfficeDocxTests(unittest.TestCase):
    def setUp(self):
        self.root = tempfile.mkdtemp(prefix="quilo-libreoffice-docx-test-")
        self.pdf_path = os.path.join(self.root, "translated.pdf")
        self.docx_path = os.path.join(self.root, "clean.docx")
        self._make_fixture()

    def tearDown(self):
        shutil.rmtree(self.root, ignore_errors=True)

    def _make_fixture(self):
        pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 80, 40), False)
        pixmap.clear_with(0x4A90E2)
        image = pixmap.tobytes("png")
        doc = fitz.open()
        for index in range(3):
            page = doc.new_page(width=595, height=842)
            page.insert_text((72, 35), "Fundamental Astronomy Korean Edition", fontsize=8)
            page.insert_text((72, 90), f"{index + 1}. Stellar Coordinates", fontsize=18)
            page.insert_text((72, 125), "Translated paragraph wraps naturally across", fontsize=11)
            page.insert_text((72, 141), "two physical lines in the source PDF.", fontsize=11)
            page.insert_text((210, 185), "E = mc^2", fontsize=14)
            page.insert_image(fitz.Rect(72, 220, 192, 280), stream=image)
            page.insert_text((72, 297), f"Figure {index + 1}. Preserved source image", fontsize=9)
            # A small, regular 2x2 lattice table.
            x0, y0, cell_w, cell_h = 72, 330, 120, 28
            for row in range(3):
                page.draw_line((x0, y0 + row * cell_h), (x0 + 2 * cell_w, y0 + row * cell_h), width=0.7)
            for col in range(3):
                page.draw_line((x0 + col * cell_w, y0), (x0 + col * cell_w, y0 + 2 * cell_h), width=0.7)
            for row, values in enumerate((("Axis", "Value"), ("RA", "12 h"))):
                for col, value in enumerate(values):
                    page.insert_text((x0 + col * cell_w + 8, y0 + row * cell_h + 18), value, fontsize=9)
            page.insert_text((292, 815), str(index + 1), fontsize=9)
        doc.save(self.pdf_path)
        doc.close()

    def test_line_wrap_normalization_and_safe_table_classification(self):
        lines = [{"text": "inter-"}, {"text": "national."}]
        self.assertEqual(MODULE._join_lines(lines), "international.")
        korean = [{"text": "별의 운동을"}, {"text": "설명한다."}]
        self.assertEqual(MODULE._join_lines(korean), "별의 운동을 설명한다.")

        block = {
            "rows": [
                {"cells": [text_cell("Axis"), text_cell("Value")]},
                {"cells": [text_cell("RA"), text_cell("12 h")]},
            ]
        }
        safe, matrix = MODULE.table_is_safe(block)
        self.assertTrue(safe)
        self.assertEqual(matrix, [["Axis", "Value"], ["RA", "12 h"]])
        block["rows"][1]["cells"][1] = text_cell("12 h", rotated=True)
        self.assertFalse(MODULE.table_is_safe(block)[0])
        formula_box = {
            "rows": [{"cells": [text_cell("E = mc^2"), text_cell("(2.30)")]}]
        }
        self.assertFalse(MODULE.table_is_safe(formula_box)[0])
        numeric_display = {
            "type": 0,
            "lines": [{
                "bbox": [0, 0, 80, 15],
                "dir": [1, 0],
                "spans": [{"bbox": [0, 0, 80, 15], "text": "8 . 79 ′′ .", "size": 11}],
            }],
        }
        self.assertEqual(MODULE._split_text_block(numeric_display, 9.0, 0)[0]["kind"], "formula")

        short_math = {
            "type": 0,
            "lines": [{
                "bbox": [0, 0, 120, 15],
                "dir": [1, 0],
                "spans": [{"bbox": [0, 0, 120, 15], "text": "sin Σ cos ( d Σ", "size": 10}],
            }],
        }
        parsed = MODULE._split_text_block(short_math, 10.0, 0)
        self.assertEqual(parsed[0]["kind"], "formula")

    def test_repairs_detached_equations_before_right_column_figures(self):
        def cell(bbox, value):
            raw = text_cell(value)
            raw["bbox"] = bbox
            raw["blocks"][0]["lines"][0]["bbox"] = bbox
            raw["blocks"][0]["lines"][0]["spans"][0]["bbox"] = bbox
            return raw

        mixed_table = {
            "bbox": (61, 208, 449, 251),
            "rows": [
                {"cells": [
                    cell((61, 208, 213, 232), "dδ = dλ sin ε cos α"),
                    cell((213, 208, 253, 251), "(2.27)"),
                    cell((253, 208, 449, 219), "그림 2.18 세차운동 설명"),
                ]},
                {"cells": [None, None, cell((253, 219, 449, 251), "장동의 작은 흔들림이다.")]},
            ],
        }
        split = MODULE._split_mixed_formula_caption_table(mixed_table, 36)
        self.assertEqual([item["kind"] for item in split], ["table_clip", "caption_clip"])
        self.assertEqual(split[0]["bbox"][2], 253.0)
        self.assertIn("장동의 작은 흔들림", split[1]["text"])

        def page(index, intro, formula_name, include_caption):
            following = [{
                "kind": "table_clip",
                "bbox": (60, 630, 250, 660),
                "page_index": index,
                "marker": formula_name,
            }]
            if include_caption:
                following.append({
                    "kind": "caption_clip",
                    "bbox": (253, 630, 450, 660),
                    "page_index": index,
                    "text": "그림 2.18 전체 캡션",
                })
            return {
                "index": index,
                "width": 504,
                "height": 720,
                "sections": [
                    {
                        "columns": [
                            {"bbox": (51, 100, 250, 620), "blocks": [{
                                "kind": "text",
                                "bbox": (53, 590, 240, 620),
                                "page_index": index,
                                "text": intro,
                            }]},
                            {"bbox": (253, 100, 454, 620), "blocks": [{
                                "kind": "image",
                                "bbox": (270, 120, 430, 300),
                                "page_index": index,
                                "image": b"image",
                            }]},
                        ]
                    },
                    {"columns": [{"bbox": (51, 630, 454, 660), "blocks": following}]},
                ],
            }

        ir = {"pages": [
            page(36, "1년 동안의 변화는 따라서 흔히", "eq-2.27", True),
            page(37, "실용적인 목적에는 근사값", "eq-2.30", False),
        ]}
        stats = MODULE._repair_detached_bottom_visuals(ir)
        self.assertEqual(stats["detached_formula_clips_reanchored"], 2)
        for target, marker in zip(ir["pages"], ("eq-2.27", "eq-2.30")):
            self.assertEqual(len(target["sections"]), 1)
            left, right = target["sections"][0]["columns"]
            self.assertEqual(left["blocks"][-1]["marker"], marker)
            self.assertEqual(right["blocks"][0]["kind"], "image")
        self.assertEqual(
            ir["pages"][0]["sections"][0]["columns"][1]["blocks"][1]["kind"],
            "caption_clip",
        )
        self.assertTrue(MODULE._is_formula_intro("실용적인 목적에는 근사값"))

    def test_groups_split_figure_captions_and_clips_math_continuations(self):
        text_figure_blocks = [
            {"kind": "caption", "bbox": (52, 53, 220, 63), "page_index": 7, "text": "그림 2.17 세차운동 때문에"},
            {"kind": "image", "bbox": (219, 56, 453, 290), "page_index": 7, "image": b"figure-a"},
            {"kind": "text", "bbox": (52, 63, 214, 124), "page_index": 7, "text": "천구의 극은 이동한다."},
            {"kind": "text", "bbox": (52, 320, 200, 340), "page_index": 7, "text": "다음 본문"},
        ]
        clipped_figure_blocks = [
            {"kind": "image", "bbox": (259, 143, 452, 262), "page_index": 9, "image": b"figure-b"},
            {"kind": "caption", "bbox": (260, 272, 432, 283), "page_index": 9, "text": "그림 2.21 망원경 설명"},
            {"kind": "formula", "bbox": (260, 283, 443, 293), "page_index": 9, "text": "t = l/c"},
            {"kind": "formula", "bbox": (260, 293, 453, 333), "page_index": 9, "text": "a = (v/c)sinθ"},
            {"kind": "text", "bbox": (260, 369, 440, 450), "page_index": 9, "text": "여기서 v는 관측자의 속도이다."},
        ]
        ir = {"pages": [{
            "sections": [{"columns": [
                {"blocks": text_figure_blocks},
                {"blocks": clipped_figure_blocks},
            ]}]
        }]}
        stats = MODULE._group_figures_and_captions(ir)
        self.assertEqual(stats, {
            "atomic_figures": 2,
            "caption_fragments_joined": 3,
            "caption_clips": 1,
        })
        first = ir["pages"][0]["sections"][0]["columns"][0]["blocks"][0]
        self.assertEqual(first["kind"], "figure")
        self.assertEqual(first["caption_mode"], "text")
        self.assertIn("천구의 극은 이동한다", first["caption_text"])
        second_column = ir["pages"][0]["sections"][0]["columns"][1]["blocks"]
        self.assertEqual(second_column[0]["kind"], "figure")
        self.assertEqual(second_column[0]["caption_mode"], "clip")
        self.assertEqual(second_column[1]["text"], "여기서 v는 관측자의 속도이다.")

    def test_builds_fresh_clean_reading_docx_with_range_and_removes_furniture(self):
        options = MODULE.BuildOptions(start_page=0, end_page=3, clip_dpi=216)
        metadata = MODULE.build_docx(self.pdf_path, self.docx_path, options)
        self.assertEqual(metadata["processed_pages"], 3)
        self.assertEqual(metadata["start_page_zero_based"], 0)
        self.assertEqual(metadata["end_page_zero_based_exclusive"], 3)
        self.assertGreaterEqual(metadata["repeating_text_removed"], 3)
        self.assertEqual(metadata["page_numbers_removed"], 3)
        self.assertGreaterEqual(metadata["headings"], 3)
        self.assertGreaterEqual(metadata["formula_clips"], 3)
        # The same source image intentionally appears on all three pages. It must
        # not be globally deduplicated merely because the bytes are identical.
        self.assertGreaterEqual(metadata["images"], 3)
        self.assertEqual(metadata["duplicate_ir_images_skipped"], 0)
        self.assertGreaterEqual(metadata["real_tables"] + metadata["table_clips"], 1)

        output = Document(self.docx_path)
        text = "\n".join(paragraph.text for paragraph in output.paragraphs)
        self.assertNotIn("Fundamental Astronomy Korean Edition", text)
        self.assertTrue(any(paragraph.style.name == "Heading 1" for paragraph in output.paragraphs))
        self.assertGreaterEqual(len(output.inline_shapes), 2)
        self.assertEqual(output.sections[0].page_width.pt, 595.0)
        self.assertEqual(output.styles["Normal"].font.name, "NanumGothic")

    def test_running_furniture_suppression_requires_page_edge_evidence(self):
        repeated_image = b"same-image-bytes"
        pages = []
        for index in range(3):
            blocks = [
                {
                    "kind": "table",
                    "bbox": (72, 350, 500, 380),
                    "page_index": index,
                    "rows": [["Chapter", str(index + 20)]],
                    "source_rows": [["Chapter", str(index + 20)]],
                },
                {
                    "kind": "image",
                    "bbox": (72, 400, 172, 460),
                    "page_index": index,
                    "image": repeated_image,
                },
            ]
            pages.append(
                {
                    "index": index,
                    "width": 595,
                    "height": 842,
                    "header": "",
                    "footer": "",
                    "sections": [{"columns": [{"bbox": (0, 0, 595, 842), "blocks": blocks}]}],
                }
            )
        ir = {"pages": pages}
        stats = MODULE._suppress_running_furniture(ir)
        self.assertEqual(stats["running_header_tables_removed"], 0)
        self.assertEqual(stats["repeating_images_removed"], 0)
        self.assertTrue(all(len(page["sections"][0]["columns"][0]["blocks"]) == 2 for page in pages))

    def test_zero_based_page_range_is_end_exclusive(self):
        options = MODULE.BuildOptions(start_page=1, end_page=3, clip_dpi=144)
        metadata = MODULE.build_docx(self.pdf_path, self.docx_path, options)
        self.assertEqual(metadata["processed_pages"], 2)
        output = Document(self.docx_path)
        text = "\n".join(paragraph.text for paragraph in output.paragraphs)
        self.assertNotIn("1. Stellar Coordinates", text)
        self.assertIn("2. Stellar Coordinates", text)
        self.assertIn("3. Stellar Coordinates", text)

    def test_rejects_invalid_page_range(self):
        with self.assertRaisesRegex(MODULE.LibreOfficeDocxError, "page range"):
            MODULE.build_docx(
                self.pdf_path,
                self.docx_path,
                MODULE.BuildOptions(start_page=2, end_page=4),
            )


if __name__ == "__main__":
    unittest.main()
